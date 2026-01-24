#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
题库转换工具 - 智能兽医大亨

支持格式：
- Word文档 (.docx)
- Excel (.xlsx, .xls)
- CSV (.csv)
- 文本文件 (.txt)
- JSON (.json)

作者: AI Assistant
版本: 2.0
"""

import pandas as pd
import json
import sys
import re
from pathlib import Path
from typing import List, Dict, Any

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  警告：未安装 python-docx，无法读取 .docx 文件")
    print("   安装方法：pip install python-docx")

class QuestionConverter:
    """题库转换器"""
    
    def __init__(self, subject_id: str = '11111111-1111-1111-1111-111111111111'):
        self.subject_id = subject_id
        self.questions = []
    
    def answer_to_index(self, answer: str) -> int:
        """将A/B/C/D转换为0/1/2/3"""
        answer = str(answer).strip().upper()
        if answer in ['A', '0', '第一个', '选项A']:
            return 0
        elif answer in ['B', '1', '第二个', '选项B']:
            return 1
        elif answer in ['C', '2', '第三个', '选项C']:
            return 2
        elif answer in ['D', '3', '第四个', '选项D']:
            return 3
        else:
            print(f"⚠️  警告：无法识别答案'{answer}'，默认使用A")
            return 0
    
    def escape_sql_string(self, text: str) -> str:
        """转义SQL字符串"""
        if pd.isna(text):
            return ""
        return str(text).replace("'", "''").strip()
    
    def from_excel(self, file_path: str) -> 'QuestionConverter':
        """从Excel/CSV读取题目"""
        print(f"📖 正在读取文件: {file_path}")
        
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8')
            else:
                df = pd.read_excel(file_path)
        except Exception as e:
            print(f"❌ 读取失败: {e}")
            print("💡 提示：请确保文件格式正确，Excel需要安装 openpyxl")
            return self
        
        print(f"✅ 成功读取 {len(df)} 行数据")
        
        # 检测列名（支持中英文）
        columns_map = {
            'question': ['题目', 'question', '问题', 'q'],
            'option_a': ['选项A', 'option_a', 'a', '选项1'],
            'option_b': ['选项B', 'option_b', 'b', '选项2'],
            'option_c': ['选项C', 'option_c', 'c', '选项3'],
            'option_d': ['选项D', 'option_d', 'd', '选项4'],
            'answer': ['答案', 'answer', 'ans', '正确答案'],
            'explanation': ['解析', 'explanation', 'exp', '说明'],
            'difficulty': ['难度', 'difficulty', 'diff', '难度等级']
        }
        
        # 智能匹配列名
        new_columns = {}
        for col in df.columns:
            col_lower = str(col).lower().strip()
            for target, variants in columns_map.items():
                if col_lower in [v.lower() for v in variants]:
                    new_columns[col] = target
                    break
        
        df.rename(columns=new_columns, inplace=True)
        
        # 验证必需列
        required = ['question', 'option_a', 'option_b', 'option_c', 'option_d', 'answer', 'explanation']
        missing = [col for col in required if col not in df.columns]
        
        if missing:
            print(f"❌ 缺少必需列: {', '.join(missing)}")
            print(f"📋 当前列名: {list(df.columns)}")
            return self
        
        # 如果没有难度列，默认为1
        if 'difficulty' not in df.columns:
            df['difficulty'] = 1
            print("💡 未找到难度列，默认设置为1")
        
        # 转换数据
        for idx, row in df.iterrows():
            try:
                question_data = {
                    'question': self.escape_sql_string(row['question']),
                    'options': [
                        self.escape_sql_string(row['option_a']),
                        self.escape_sql_string(row['option_b']),
                        self.escape_sql_string(row['option_c']),
                        self.escape_sql_string(row['option_d'])
                    ],
                    'correct_answer': self.answer_to_index(row['answer']),
                    'explanation': self.escape_sql_string(row['explanation']),
                    'difficulty': int(row.get('difficulty', 1))
                }
                
                # 验证难度范围
                if question_data['difficulty'] < 1 or question_data['difficulty'] > 5:
                    question_data['difficulty'] = 1
                
                self.questions.append(question_data)
            except Exception as e:
                print(f"⚠️  第 {idx+2} 行数据解析失败: {e}")
                continue
        
        print(f"✅ 成功解析 {len(self.questions)} 道题目")
        return self
    
    def from_text(self, file_path: str) -> 'QuestionConverter':
        """从文本文件读取题目（支持多种格式）"""
        print(f"📖 正在读取文本文件: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 尝试解析格式1：标准格式
        # 1. 题目？
        # A. 选项1
        # B. 选项2
        # C. 选项3
        # D. 选项4
        # 答案：B
        # 解析：...
        
        pattern = r'(?:\d+[\.\、])\s*(.*?)\n\s*[AaＡａ][\.\、、]\s*(.*?)\n\s*[BbＢｂ][\.\、、]\s*(.*?)\n\s*[CcＣｃ][\.\、、]\s*(.*?)\n\s*[DdＤｄ][\.\、、]\s*(.*?)\n\s*(?:答案|正确答案)[:：]\s*([A-Da-d])\n\s*(?:解析|说明)[:：]\s*(.*?)(?=\n\d+[\.\、]|\Z)'
        
        matches = re.findall(pattern, content, re.DOTALL)
        
        for match in matches:
            try:
                question_data = {
                    'question': self.escape_sql_string(match[0]),
                    'options': [
                        self.escape_sql_string(match[1]),
                        self.escape_sql_string(match[2]),
                        self.escape_sql_string(match[3]),
                        self.escape_sql_string(match[4])
                    ],
                    'correct_answer': self.answer_to_index(match[5]),
                    'explanation': self.escape_sql_string(match[6]),
                    'difficulty': 1  # 默认难度
                }
                self.questions.append(question_data)
            except Exception as e:
                print(f"⚠️  解析题目失败: {e}")
                continue
        
        if len(self.questions) == 0:
            print("❌ 未能解析出题目，请检查文件格式")
            print("📝 标准格式示例：")
            print("""
1. 犬的正常体温范围是？
A. 37.5-38.5°C
B. 38.0-39.2°C
C. 39.5-40.5°C
D. 36.0-37.5°C
答案：B
解析：犬的正常体温为38.0-39.2°C...
            """)
        else:
            print(f"✅ 成功解析 {len(self.questions)} 道题目")
        
        return self
    
    def from_docx(self, file_path: str) -> 'QuestionConverter':
        """从Word文档读取题目"""
        if not DOCX_AVAILABLE:
            print("❌ 请先安装 python-docx: pip install python-docx")
            return self
        
        print(f"📖 正在读取Word文档: {file_path}")
        
        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"❌ 读取Word文档失败: {e}")
            return self
        
        # 提取所有段落文本
        full_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        
        print(f"📄 文档总字数: {len(full_text)}")
        
        # 尝试多种格式解析
        
        # 格式1：表格格式（Word中的表格）
        if len(doc.tables) > 0:
            print("🔍 检测到表格，尝试从表格读取...")
            questions_from_table = self._parse_word_table(doc.tables)
            if questions_from_table > 0:
                print(f"✅ 从表格成功解析 {questions_from_table} 道题目")
                return self
        
        # 格式2：标准文本格式
        print("🔍 尝试从文本内容解析...")
        
        # 模式1: 标准格式
        # 1. 题目？
        # A. 选项1
        # B. 选项2
        # C. 选项3
        # D. 选项4
        # 答案：B
        # 解析：...
        
        pattern1 = r'(?:\d+[\.\、、])\s*(.*?)\n\s*[AaＡａ][\.\、、、]\s*(.*?)\n\s*[BbＢｂ][\.\、、、]\s*(.*?)\n\s*[CcＣｃ][\.\、、、]\s*(.*?)\n\s*[DdＤｄ][\.\、、、]\s*(.*?)\n\s*(?:答案|正确答案|参考答案)[:：]\s*([A-Da-d])\n\s*(?:解析|说明|详解)[:：]\s*(.*?)(?=\n\d+[\.\、、]|\Z)'
        
        matches = re.findall(pattern1, full_text, re.DOTALL)
        
        if matches:
            for match in matches:
                try:
                    question_data = {
                        'question': self.escape_sql_string(match[0]),
                        'options': [
                            self.escape_sql_string(match[1]),
                            self.escape_sql_string(match[2]),
                            self.escape_sql_string(match[3]),
                            self.escape_sql_string(match[4])
                        ],
                        'correct_answer': self.answer_to_index(match[5]),
                        'explanation': self.escape_sql_string(match[6]),
                        'difficulty': 1
                    }
                    self.questions.append(question_data)
                except Exception as e:
                    print(f"⚠️  解析题目失败: {e}")
                    continue
        
        # 模式2: 简化格式（无解析）
        if len(self.questions) == 0:
            pattern2 = r'(?:\d+[\.\、、])\s*(.*?)\n\s*[AaＡａ][\.\、、、]\s*(.*?)\n\s*[BbＢｂ][\.\、、、]\s*(.*?)\n\s*[CcＣｃ][\.\、、、]\s*(.*?)\n\s*[DdＤｄ][\.\、、、]\s*(.*?)\n\s*(?:答案|正确答案)[:：]\s*([A-Da-d])'
            
            matches = re.findall(pattern2, full_text, re.DOTALL)
            
            for match in matches:
                try:
                    question_data = {
                        'question': self.escape_sql_string(match[0]),
                        'options': [
                            self.escape_sql_string(match[1]),
                            self.escape_sql_string(match[2]),
                            self.escape_sql_string(match[3]),
                            self.escape_sql_string(match[4])
                        ],
                        'correct_answer': self.answer_to_index(match[5]),
                        'explanation': '暂无解析',
                        'difficulty': 1
                    }
                    self.questions.append(question_data)
                except Exception as e:
                    print(f"⚠️  解析题目失败: {e}")
                    continue
        
        if len(self.questions) == 0:
            print("❌ 未能解析出题目，请检查Word文档格式")
            print("\n📝 支持的格式：")
            print("""
方式1 - 表格格式：
┌────────┬────────┬────────┬────────┬────────┬──────┬──────┬────┐
│ 题目   │ 选项A  │ 选项B  │ 选项C  │ 选项D  │ 答案 │ 解析 │难度│
├────────┼────────┼────────┼────────┼────────┼──────┼──────┼────┤
│ 题目1？│ 选项1  │ 选项2  │ 选项3  │ 选项4  │  B   │ ...  │ 1  │
└────────┴────────┴────────┴────────┴────────┴──────┴──────┴────┘

方式2 - 文本格式：
1. 犬的正常体温范围是？
A. 37.5-38.5°C
B. 38.0-39.2°C
C. 39.5-40.5°C
D. 36.0-37.5°C
答案：B
解析：犬的正常体温为38.0-39.2°C...

2. 下一题...
            """)
        else:
            print(f"✅ 成功解析 {len(self.questions)} 道题目")
        
        return self
    
    def _parse_word_table(self, tables) -> int:
        """解析Word文档中的表格"""
        count = 0
        
        for table in tables:
            # 假设第一行是表头
            if len(table.rows) < 2:
                continue
            
            # 获取表头，确定列的位置
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            
            # 智能匹配列
            col_map = {}
            for idx, header in enumerate(header_cells):
                if any(x in header for x in ['题目', 'question', '问题']):
                    col_map['question'] = idx
                elif any(x in header for x in ['选项a', 'option_a', 'a', '选项1']):
                    col_map['option_a'] = idx
                elif any(x in header for x in ['选项b', 'option_b', 'b', '选项2']):
                    col_map['option_b'] = idx
                elif any(x in header for x in ['选项c', 'option_c', 'c', '选项3']):
                    col_map['option_c'] = idx
                elif any(x in header for x in ['选项d', 'option_d', 'd', '选项4']):
                    col_map['option_d'] = idx
                elif any(x in header for x in ['答案', 'answer', 'ans']):
                    col_map['answer'] = idx
                elif any(x in header for x in ['解析', 'explanation', '说明']):
                    col_map['explanation'] = idx
                elif any(x in header for x in ['难度', 'difficulty', 'diff']):
                    col_map['difficulty'] = idx
            
            # 检查必需列
            required = ['question', 'option_a', 'option_b', 'option_c', 'option_d', 'answer']
            if not all(k in col_map for k in required):
                continue
            
            # 解析每一行
            for row in table.rows[1:]:  # 跳过表头
                try:
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    if len(cells) <= max(col_map.values()):
                        continue
                    
                    question_data = {
                        'question': self.escape_sql_string(cells[col_map['question']]),
                        'options': [
                            self.escape_sql_string(cells[col_map['option_a']]),
                            self.escape_sql_string(cells[col_map['option_b']]),
                            self.escape_sql_string(cells[col_map['option_c']]),
                            self.escape_sql_string(cells[col_map['option_d']])
                        ],
                        'correct_answer': self.answer_to_index(cells[col_map['answer']]),
                        'explanation': self.escape_sql_string(
                            cells[col_map.get('explanation', col_map['question'])]
                        ) if 'explanation' in col_map else '暂无解析',
                        'difficulty': int(cells[col_map['difficulty']]) if 'difficulty' in col_map and cells[col_map['difficulty']].isdigit() else 1
                    }
                    
                    # 验证数据完整性
                    if question_data['question'] and all(question_data['options']):
                        self.questions.append(question_data)
                        count += 1
                except Exception as e:
                    print(f"⚠️  解析表格行失败: {e}")
                    continue
        
        return count
    
    def from_json(self, file_path: str) -> 'QuestionConverter':
        """从JSON文件读取题目"""
        print(f"📖 正在读取JSON文件: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if isinstance(data, list):
            for item in data:
                try:
                    question_data = {
                        'question': self.escape_sql_string(item.get('question', '')),
                        'options': [
                            self.escape_sql_string(opt) 
                            for opt in item.get('options', ['', '', '', ''])[:4]
                        ],
                        'correct_answer': item.get('correct_answer', 0),
                        'explanation': self.escape_sql_string(item.get('explanation', '')),
                        'difficulty': int(item.get('difficulty', 1))
                    }
                    self.questions.append(question_data)
                except Exception as e:
                    print(f"⚠️  解析题目失败: {e}")
                    continue
        
        print(f"✅ 成功解析 {len(self.questions)} 道题目")
        return self
    
    def to_sql(self, output_file: str = None) -> str:
        """生成SQL语句"""
        if not self.questions:
            print("❌ 没有题目可以转换")
            return ""
        
        print(f"\n🔄 正在生成SQL语句...")
        
        sql_lines = [
            "-- 自动生成的题目导入SQL",
            f"-- 科目ID: {self.subject_id}",
            f"-- 题目数量: {len(self.questions)}",
            f"-- 生成时间: {pd.Timestamp.now()}",
            "",
            "INSERT INTO training_questions (subject_id, question, options, correct_answer, explanation, difficulty) VALUES"
        ]
        
        values = []
        for q in self.questions:
            options_json = json.dumps(q['options'], ensure_ascii=False)
            
            value = f"""('{self.subject_id}',
 '{q['question']}',
 '{options_json}',
 {q['correct_answer']},
 '{q['explanation']}',
 {q['difficulty']})"""
            
            values.append(value)
        
        sql_lines.append(",\n\n".join(values))
        sql_lines.append("\nON CONFLICT DO NOTHING;")
        
        sql_content = '\n'.join(sql_lines)
        
        # 写入文件
        if output_file is None:
            output_file = f"imported_questions_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.sql"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(sql_content)
        
        print(f"✅ SQL文件已生成: {output_file}")
        print(f"📊 包含 {len(self.questions)} 道题目")
        
        return sql_content
    
    def preview(self, count: int = 3):
        """预览前几道题目"""
        if not self.questions:
            print("❌ 没有题目可以预览")
            return
        
        print(f"\n📋 预览前 {min(count, len(self.questions))} 道题目：\n")
        print("=" * 80)
        
        for i, q in enumerate(self.questions[:count], 1):
            print(f"\n第 {i} 题：{q['question']}")
            for j, opt in enumerate(q['options']):
                marker = "✓" if j == q['correct_answer'] else " "
                print(f"  {chr(65+j)}. {opt} {marker}")
            print(f"  解析：{q['explanation']}")
            print(f"  难度：{'⭐' * q['difficulty']}")
            print("-" * 80)


def main():
    """主函数"""
    print("""
╔══════════════════════════════════════════════════════════╗
║          题库转换工具 - 智能兽医大亨                      ║
║                    Version 1.0                           ║
╚══════════════════════════════════════════════════════════╝
    """)
    
    if len(sys.argv) < 2:
        print("📖 用法：")
        print("  python questions_converter.py <输入文件> [科目ID] [输出文件]")
        print("\n📝 示例：")
        print("  python questions_converter.py questions.docx")
        print("  python questions_converter.py questions.xlsx")
        print("  python questions_converter.py questions.csv 科目ID-123")
        print("  python questions_converter.py questions.txt")
        print("\n支持格式：.docx, .xlsx, .xls, .csv, .txt, .json")
        sys.exit(1)
    
    input_file = sys.argv[1]
    subject_id = sys.argv[2] if len(sys.argv) > 2 else '11111111-1111-1111-1111-111111111111'
    output_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not Path(input_file).exists():
        print(f"❌ 文件不存在: {input_file}")
        sys.exit(1)
    
    # 创建转换器
    converter = QuestionConverter(subject_id)
    
    # 根据文件类型选择读取方法
    file_ext = Path(input_file).suffix.lower()
    
    if file_ext == '.docx':
        converter.from_docx(input_file)
    elif file_ext in ['.xlsx', '.xls', '.csv']:
        converter.from_excel(input_file)
    elif file_ext == '.txt':
        converter.from_text(input_file)
    elif file_ext == '.json':
        converter.from_json(input_file)
    else:
        print(f"❌ 不支持的文件类型: {file_ext}")
        print(f"   支持格式: .docx, .xlsx, .xls, .csv, .txt, .json")
        sys.exit(1)
    
    # 预览
    converter.preview(3)
    
    # 生成SQL
    if converter.questions:
        print("\n" + "=" * 80)
        converter.to_sql(output_file)
        print("\n✨ 完成！下一步：在 Supabase SQL Editor 中执行生成的SQL文件")
    else:
        print("\n❌ 转换失败，请检查文件格式")


if __name__ == '__main__':
    main()
