#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 题库导入工具 - 将 Word 文档转换为 Supabase SQL

支持格式：
1. 共用题干 + 多选题
2. 普通单选/多选题
3. 自动答案识别

依赖：python-docx
安装：pip install python-docx

作者: AI Assistant
版本: 1.0
"""

import re
import json
from pathlib import Path
import sys

try:
    from docx import Document
except ImportError:
    print("❌ 缺少 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)


def extract_text_from_docx(docx_path: str) -> str:
    """从 docx 文件提取文本"""
    try:
        doc = Document(docx_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # 从表格中提取文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        full_text = '\n'.join(text_parts)
        print(f"✅ 成功提取 {len(full_text)} 个字符")
        return full_text
    except Exception as e:
        print(f"❌ 读取 docx 文件失败: {e}")
        sys.exit(1)


def escape_sql_string(text: str) -> str:
    """转义 SQL 字符串"""
    if not text:
        return ""
    return str(text).replace("'", "''").strip()


def answer_to_index(answer: str) -> int:
    """将 A/B/C/D/E 转换为 0-4"""
    answer = str(answer).strip().upper()
    mapping = {'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4}
    return mapping.get(answer, 0)


def parse_questions_from_text(text: str):
    """智能解析题目"""
    questions = []
    
    # 检测是否有共用题干
    has_case_stem = bool(re.search(r'[（\(](\d+)[～~\-\-](\d+)题?共用题干[）\)]', text))
    
    if has_case_stem:
        print("🔍 检测到共用题干格式")
        questions = parse_case_based_text(text)
    else:
        print("🔍 检测到普通题目格式")
        questions = parse_normal_text(text)
    
    return questions


def parse_case_based_text(text):
    """解析共用题干格式"""
    questions = []
    
    # 分割不同的案例组
    case_splits = re.split(r'[（\(](\d+)[～~\-\-](\d+)题?共用题干[）\)]\s*', text)
    
    i = 0
    while i < len(case_splits):
        if i == 0:
            i += 1
            continue
        
        if i + 2 < len(case_splits):
            start_num = case_splits[i]
            end_num = case_splits[i+1]
            content = case_splits[i+2]
            
            print(f"📖 处理题目 {start_num}-{end_num}")
            
            # 提取题干
            first_q_match = re.search(r'(\d+)\s*[\.\.、]', content)
            if first_q_match:
                stem_end = first_q_match.start()
                case_stem = content[:stem_end].strip()
                questions_part = content[stem_end:]
                
                # 分割成单个题目
                q_blocks = re.split(r'(?=\d+\s*[\.\.、])', questions_part)
                
                for block in q_blocks:
                    if not block.strip():
                        continue
                    
                    q = parse_single_question(block, case_stem)
                    if q:
                        questions.append(q)
            
            i += 3
        else:
            i += 1
    
    return questions


def parse_normal_text(text):
    """解析普通题目格式"""
    questions = []
    
    # 分割成单个题目
    q_blocks = re.split(r'(?=\d+\s*[\.\.、])', text)
    
    for block in q_blocks:
        if not block.strip():
            continue
        
        q = parse_single_question(block)
        if q:
            questions.append(q)
    
    return questions


def parse_single_question(block, stem=""):
    """解析单个题目块"""
    lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
    
    if len(lines) < 2:
        return None
    
    # 提取题号和问题
    first_line = lines[0]
    q_match = re.match(r'(\d+)\s*[\.\.、]\s*(.*)', first_line)
    if not q_match:
        return None
    
    question_num = q_match.group(1)
    question_text = q_match.group(2).strip()
    
    # 合并题干
    if stem:
        full_question = stem + " " + question_text
    else:
        full_question = question_text
    
    # 提取选项
    options = []
    answer = None
    explanation = ""
    subject = "综合"  # 默认科目
    
    for line in lines[1:]:
        # 匹配选项 A. B. C. D. E.
        opt_match = re.match(r'([A-Ea-e])[\.\.、]\s*(.*)', line)
        if opt_match:
            options.append(opt_match.group(2).strip())
            continue
        
        # 匹配答案
        ans_match = re.match(r'(?:答案|正确答案|参考答案)[:：]\s*([A-Ea-e\d]+)', line, re.IGNORECASE)
        if ans_match:
            answer = ans_match.group(1)
            continue
        
        # 匹配解析
        exp_match = re.match(r'(?:解析|说明|详解|分析)[:：]\s*(.*)', line, re.IGNORECASE)
        if exp_match:
            explanation = exp_match.group(1).strip()
            continue
        
        # 匹配科目
        subj_match = re.match(r'(?:科目|类别)[:：]\s*(.*)', line, re.IGNORECASE)
        if subj_match:
            subject = subj_match.group(1).strip()
            continue
    
    # 处理选项数量
    if len(options) < 4:
        print(f"⚠️  题目 {question_num} 选项不足4个，跳过")
        return None
    
    # 如果有5个选项，只取前4个
    if len(options) > 4:
        options = options[:4]
    
    # 处理答案
    correct_answer = 0
    if answer:
        correct_answer = answer_to_index(answer)
    else:
        print(f"⚠️  题目 {question_num} 未找到答案，默认使用A")
    
    # 处理解析
    if not explanation:
        explanation = "暂无解析"
    
    return {
        'number': question_num,
        'question': full_question,
        'options': options,
        'correct_answer': correct_answer,
        'explanation': explanation,
        'subject': subject,
        'difficulty': 2,
        'question_type': 'single_choice'
    }


def generate_sql(questions, subject='综合应用', output_file=None):
    """生成 Supabase SQL 语句 - 用于 vet_exam_questions 表"""
    if not questions:
        print("❌ 没有题目可以转换")
        return ""
    
    # 科目映射 - 将缩写或其他名称映射为标准科目名称
    subject_mapping = {
        '基础': '基础兽医学',
        '预防': '预防兽医学',
        '临床': '临床兽医学',
        '综合': '综合应用',
        '兽医法规': '基础兽医学',  # 法规归入基础
        '兽医内科学': '临床兽医学',
        '兽医外科学': '临床兽医学',
        '兽医传染病学': '预防兽医学',
        '兽医药理学': '基础兽医学',
    }
    
    # 映射科目名称
    mapped_subject = subject_mapping.get(subject, subject)
    
    # 验证科目是否为四个标准科目之一
    valid_subjects = ['基础兽医学', '预防兽医学', '临床兽医学', '综合应用']
    if mapped_subject not in valid_subjects:
        print(f"⚠️  科目 '{subject}' 映射为 '{mapped_subject}'，将使用默认科目'综合应用'")
        mapped_subject = '综合应用'
    else:
        print(f"✅ 科目：{mapped_subject}")
    
    sql_lines = [
        "-- DOCX 转换生成的执业兽医考试题目 SQL",
        f"-- 科目: {mapped_subject}",
        f"-- 题目数量: {len(questions)}",
        "",
        "-- 插入到执业兽医考试题库",
        "INSERT INTO vet_exam_questions (question_type, stem, is_shared_stem, options, correct_answer, explanation, subject, difficulty, is_real_exam) VALUES"
    ]
    
    values = []
    for q in questions:
        options_json = json.dumps(q['options'], ensure_ascii=False)
        
        value = f"""('single',
 '{escape_sql_string(q['question'])}',
 false,
 '{options_json}'::jsonb,
 '{q['correct_answer']}'::jsonb,
 '{escape_sql_string(q['explanation'])}',
 '{mapped_subject}',
 {q['difficulty']},
 true)"""
        
        values.append(value)
    
    sql_lines.append(",\n\n".join(values))
    sql_lines.append("\nON CONFLICT DO NOTHING;")
    
    sql_content = '\n'.join(sql_lines)
    
    # 写入文件
    if output_file is None:
        output_file = "docx_converted_questions.sql"
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(sql_content)
    
    print(f"\n✅ SQL文件已生成: {output_file}")
    print(f"📊 包含 {len(questions)} 道题目")
    print(f"📁 目标表: vet_exam_questions")
    print(f"📚 科目分类: {mapped_subject}")
    
    return sql_content



def main():
    print("""
╔══════════════════════════════════════════════════════════╗
║          DOCX 题库导入工具 v1.0                           ║
║          将 Word 文档转换为 Supabase SQL                 ║
╚══════════════════════════════════════════════════════════╝
    """)
    
    if len(sys.argv) < 2:
        print("📖 用法：")
        print("  python docx_to_sql.py <docx文件> [科目] [输出文件]")
        print("\n📝 示例：")
        print("  python docx_to_sql.py questions.docx")
        print("  python docx_to_sql.py questions.docx 兽医内科学")
        print("  python docx_to_sql.py questions.docx 兽医外科学 output.sql")
        sys.exit(1)
    
    # 解析参数
    docx_file = sys.argv[1]
    subject = sys.argv[2] if len(sys.argv) > 2 else '综合'
    output_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not Path(docx_file).exists():
        print(f"❌ 文件不存在: {docx_file}")
        sys.exit(1)
    
    # 提取文本
    print(f"📖 正在读取 DOCX 文件: {docx_file}")
    text = extract_text_from_docx(docx_file)
    
    # 解析题目
    print("\n🔍 正在智能解析题目...")
    questions = parse_questions_from_text(text)
    
    if not questions:
        print("❌ 未能解析出题目，请检查文件格式")
        sys.exit(1)
    
    # 预览
    print(f"\n📋 预览前 3 道题目：")
    print("=" * 80)
    for i, q in enumerate(questions[:3], 1):
        print(f"\n题目 {i}：{q['question'][:60]}...")
        for j, opt in enumerate(q['options']):
            marker = "✓" if j == q['correct_answer'] else " "
            print(f"  {chr(65+j)}. {opt[:40]}... {marker}")
        print(f"  答案：{chr(65+q['correct_answer'])}")
        print(f"  解析：{q['explanation'][:50]}...")
        print("-" * 80)
    
    # 生成SQL
    print("\n" + "=" * 80)
    generate_sql(questions, subject, output_file)
    print("\n✨ 完成！")
    print("\n📌 下一步：")
    print("  1. 打开 Supabase SQL Editor")
    print("  2. 复制生成的 SQL 文件内容")
    print("  3. 粘贴并执行 SQL")
    print("  4. 刷新应用即可看到新题目")


if __name__ == '__main__':
    main()
